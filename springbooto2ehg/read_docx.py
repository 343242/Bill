from docx import Document

# 读取docx文件
doc = Document('f:/javapro/bs/springbooto2ehg/需求.docx')

# 遍历所有段落并打印内容
for para in doc.paragraphs:
    text = para.text
    if text:
        print(text)

# 遍历所有表格并打印内容
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if text:
                print(text)